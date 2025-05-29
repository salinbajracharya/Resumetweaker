from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
import os
from dotenv import load_dotenv
import tempfile
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import json
from sentence_transformers import SentenceTransformer, util
from transformers import pipeline
import torch

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Initialize models (load only once at startup)
@app.before_first_request
def load_models():
    global sentence_model, keyword_extractor
    # Load sentence transformer model for similarity
    sentence_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    # Load keyword extraction model
    keyword_extractor = pipeline("token-classification", 
                               model="yanekyuk/bert-uncased-keyword-extractor")

# Add JSON template filter
@app.template_filter('from_json')
def from_json(value):
    return json.loads(value)

# Configuration
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///itsutra.db')
if app.config['SQLALCHEMY_DATABASE_URI'].startswith('postgres://'):
    app.config['SQLALCHEMY_DATABASE_URI'] = app.config['SQLALCHEMY_DATABASE_URI'].replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'info'

# User model
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    first_name = db.Column(db.String(64))
    last_name = db.Column(db.String(64))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_active = db.Column(db.Boolean, default=True)
    last_login = db.Column(db.DateTime)
    # Add relationship to ResumeHistory
    resume_history = db.relationship('ResumeHistory', backref='user', lazy=True)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    @property
    def full_name(self):
        return f"{self.first_name} {self.last_name}"

class ResumeHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    job_title = db.Column(db.String(200), nullable=False)
    job_description = db.Column(db.Text, nullable=False)
    resume_text = db.Column(db.Text, nullable=False)
    analysis_result = db.Column(db.Text, nullable=False)  # Stored as JSON
    match_score = db.Column(db.Float, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    domain_compatibility = db.Column(db.String(50), nullable=False)
    
    def to_dict(self):
        return {
            'id': self.id,
            'job_title': self.job_title,
            'match_score': self.match_score,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'domain_compatibility': self.domain_compatibility
        }

@login_manager.user_loader
def load_user(id):
    return User.query.get(int(id))

def is_valid_email(email):
    domain = email.split('@')[-1]
    return domain.lower() == 'itsutra.com'

# Routes
@app.route('/')
@login_required
def index():
    return render_template('main/index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        remember = request.form.get('remember', False)

        if not email or not password:
            flash('Please enter both email and password', 'error')
            return redirect(url_for('login'))

        if not is_valid_email(email):
            flash('Please use your ITSutra email address', 'error')
            return redirect(url_for('login'))

        user = User.query.filter_by(email=email).first()
        if user is None or not user.check_password(password):
            flash('Invalid email or password', 'error')
            return redirect(url_for('login'))

        if not user.is_active:
            flash('Your account is inactive. Please contact admin.', 'error')
            return redirect(url_for('login'))

        login_user(user, remember=remember)
        user.last_login = datetime.utcnow()
        db.session.commit()

        next_page = request.args.get('next')
        if not next_page or next_page == '/':
            next_page = url_for('index')
        return redirect(next_page)

    return render_template('auth/login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')

        if not all([email, password, first_name, last_name]):
            flash('Please fill in all fields', 'error')
            return redirect(url_for('register'))

        if not is_valid_email(email):
            flash('Please use your ITSutra email address', 'error')
            return redirect(url_for('register'))

        if User.query.filter_by(email=email).first():
            flash('Email already registered', 'error')
            return redirect(url_for('register'))

        user = User(
            email=email,
            first_name=first_name,
            last_name=last_name
        )
        user.set_password(password)
        
        db.session.add(user)
        db.session.commit()

        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('auth/register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/profile')
@login_required
def profile():
    return render_template('main/profile.html')

@app.route('/history')
@login_required
def history():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    history = ResumeHistory.query.filter_by(user_id=current_user.id)\
        .order_by(ResumeHistory.created_at.desc())\
        .paginate(page=page, per_page=per_page)
    return render_template('main/history.html', history=history)

@app.route('/history/<int:history_id>')
@login_required
def view_history(history_id):
    entry = ResumeHistory.query.get_or_404(history_id)
    if entry.user_id != current_user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('history'))
    return render_template('main/history_detail.html', entry=entry)

def extract_text_from_docx(file_path):
    """Extract text from a .docx file."""
    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

def extract_text_from_txt(file_path):
    """Extract text from a .txt file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_uploaded_file(file):
    """Extract text from an uploaded file."""
    if not file:
        return ""
        
    # Save the uploaded file temporarily
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, file.filename)
    file.save(temp_path)
    
    try:
        # Extract text based on file type
        if file.filename.lower().endswith('.docx'):
            text = extract_text_from_docx(temp_path)
        elif file.filename.lower().endswith('.txt'):
            text = extract_text_from_txt(temp_path)
        elif file.filename.lower().endswith('.pdf'):
            # For now, we'll just return a message
            text = "PDF support coming soon"
        else:
            text = "Unsupported file format"
    finally:
        # Clean up
        os.remove(temp_path)
        os.rmdir(temp_dir)
    
    return text

def analyze_resume(resume_text, jd_text, job_title):
    """Analyze resume against job description with domain compatibility check."""
    try:
        # Encode texts for similarity comparison
        resume_embedding = sentence_model.encode(resume_text, convert_to_tensor=True)
        jd_embedding = sentence_model.encode(jd_text, convert_to_tensor=True)
        
        # Calculate similarity score
        similarity = float(util.pytorch_cos_sim(resume_embedding, jd_embedding)[0][0])
        
        # Extract keywords from job description
        keywords = keyword_extractor(jd_text)
        # Filter and sort keywords by score
        unique_keywords = {}
        for kw in keywords:
            if kw['entity'] == 'KEY':
                word = kw['word'].strip()
                if word not in unique_keywords or kw['score'] > unique_keywords[word]:
                    unique_keywords[word] = kw['score']
        
        # Sort keywords by score
        sorted_keywords = sorted(unique_keywords.items(), key=lambda x: x[1], reverse=True)[:10]
        keywords_list = [k[0] for k in sorted_keywords]
        
        # Check which keywords are missing from resume
        missing_keywords = [
            keyword for keyword in keywords_list 
            if keyword.lower() not in resume_text.lower()
        ]
        
        # Extract domain from job title and resume
        job_domains = extract_domains(job_title, jd_text)
        resume_domains = extract_domains_from_resume(resume_text)
        
        # Check domain compatibility
        domain_match = check_domain_compatibility(job_domains, resume_domains)
        
        return {
            'similarity': similarity,
            'keywords': keywords_list,
            'missing_keywords': missing_keywords,
            'domain_compatibility': domain_match['status'],
            'domain_message': domain_match['message'],
            'job_domains': job_domains,
            'resume_domains': resume_domains
        }
    except Exception as e:
        print(f"Analysis error: {str(e)}")
        return {
            'similarity': 0,
            'keywords': [],
            'missing_keywords': [],
            'domain_compatibility': 'error',
            'domain_message': 'Error analyzing resume',
            'job_domains': [],
            'resume_domains': []
        }

def extract_domains(job_title, jd_text):
    """Extract potential domains from job title and description."""
    domains = {
        'automotive': ['automotive', 'vehicle', 'car', 'oem'],
        'healthcare': ['healthcare', 'medical', 'hospital', 'clinical'],
        'finance': ['finance', 'banking', 'investment', 'trading'],
        'technology': ['software', 'web', 'cloud', 'api', 'development'],
        'manufacturing': ['manufacturing', 'production', 'industrial'],
        # Add more domains as needed
    }
    
    found_domains = set()
    text = f"{job_title.lower()} {jd_text.lower()}"
    
    for domain, keywords in domains.items():
        if any(keyword in text for keyword in keywords):
            found_domains.add(domain)
    
    return list(found_domains)

def extract_domains_from_resume(resume_text):
    """Extract domains from resume based on experience and keywords."""
    return extract_domains("", resume_text)

def check_domain_compatibility(job_domains, resume_domains):
    """Check if the resume domains match the job domains."""
    common_domains = set(job_domains) & set(resume_domains)
    
    if not job_domains:
        return {
            'status': 'neutral',
            'message': 'No specific domain requirements identified.'
        }
    
    if not common_domains:
        return {
            'status': 'incompatible',
            'message': f'Your experience does not match the required domain(s): {", ".join(job_domains)}. Consider highlighting any relevant transferable skills.'
        }
    
    if len(common_domains) == len(job_domains):
        return {
            'status': 'compatible',
            'message': 'Your experience aligns well with the required domain(s).'
        }
    
    return {
        'status': 'partial',
        'message': f'You have experience in some required domains: {", ".join(common_domains)}. Consider emphasizing these areas.'
    }

def create_analysis_document(resume_text, jd_text, analysis):
    """Create a document with resume analysis."""
    doc = docx.Document()
    
    # Add header
    header = doc.add_heading('Resume Analysis Report', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add similarity score
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Match Score: ').bold = True
    p.add_run(f"{analysis['similarity']*100:.1f}%")
    
    # Add key findings section
    doc.add_heading('Key Findings', level=1)
    
    # Important keywords section
    doc.add_heading('Important Keywords from Job Description:', level=2)
    if analysis['keywords']:
        for keyword in analysis['keywords']:
            doc.add_paragraph(f"• {keyword}", style='List Bullet')
    else:
        doc.add_paragraph("No keywords extracted")
    
    # Missing keywords section
    doc.add_heading('Missing Keywords:', level=2)
    if analysis['missing_keywords']:
        for keyword in analysis['missing_keywords']:
            doc.add_paragraph(f"• {keyword}", style='List Bullet')
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Recommendation: ').bold = True
        p.add_run('Consider incorporating these missing keywords into your resume where applicable.')
    else:
        doc.add_paragraph("Your resume contains all the important keywords!")
    
    return doc

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    if request.method == 'POST':
        if 'resume' not in request.files:
            flash('No resume file uploaded', 'error')
            return redirect(request.url)
        
        resume = request.files['resume']
        if resume.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        
        # Get job title
        job_title = request.form.get('job_title', '').strip()
        if not job_title:
            flash('Please provide a job title', 'error')
            return redirect(request.url)
        
        # Get job description (either from file or text)
        jd_input_type = request.form.get('jd_input_type')
        jd_text = ''
        
        if jd_input_type == 'file':
            if 'jd_file' not in request.files:
                flash('No job description file uploaded', 'error')
                return redirect(request.url)
            
            jd_file = request.files['jd_file']
            if jd_file.filename == '':
                flash('No job description file selected', 'error')
                return redirect(request.url)
            
            jd_text = extract_text_from_uploaded_file(jd_file)
        else:
            jd_text = request.form.get('jd_text', '').strip()
            if not jd_text:
                flash('Please provide a job description', 'error')
                return redirect(request.url)
        
        # Process the resume
        resume_text = extract_text_from_uploaded_file(resume)
        
        # Analyze resume
        analysis = analyze_resume(resume_text, jd_text, job_title)
        
        # Save to history
        history_entry = ResumeHistory(
            user_id=current_user.id,
            job_title=job_title,
            job_description=jd_text,
            resume_text=resume_text,
            analysis_result=json.dumps(analysis),
            match_score=analysis['similarity'],
            domain_compatibility=analysis['domain_compatibility']
        )
        db.session.add(history_entry)
        db.session.commit()
        
        # Create analysis document
        doc = create_analysis_document(resume_text, jd_text, analysis)
        
        # Save to temporary file and send
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, 'resume_analysis.docx')
        doc.save(output_path)
        
        try:
            return send_file(
                output_path,
                as_attachment=True,
                download_name='resume_analysis.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        finally:
            os.remove(output_path)
            os.rmdir(temp_dir)
    
    return render_template('upload.html')

# Context processor for template variables
@app.context_processor
def utility_processor():
    return {'now': datetime.utcnow()}

# Create database tables
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(debug=True) 